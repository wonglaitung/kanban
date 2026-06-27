"""
Kanban AI 服务

提供任务数据字典和查询接口，供 AI 理解和查询任务数据。
参考: /data/bank-services-plugins/docs/backend_api_spec.md

API 设计:
1. GET /api/ai/dictionary - 返回任务字段描述
2. GET /api/ai/query - 查询任务数据
3. POST /api/ai/chat - AI 对话接口（使用 Harness SDK）
"""

import asyncio
import json
import os
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional

from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Query
from pydantic import BaseModel

# 加载根目录的 .env 文件
ROOT_DIR = Path(__file__).parent.parent
load_dotenv(ROOT_DIR / ".env")
load_dotenv(ROOT_DIR / ".env.local")
load_dotenv(ROOT_DIR / ".env.production")

# 添加 SDK 路径
SDK_PATH = Path(__file__).parent.parent.parent / "harness" / "packages" / "sdk" / "src"
import sys

if SDK_PATH.exists():
    sys.path.insert(0, str(SDK_PATH))

from config.dictionary import (
    ALLOWED_FILTERS,
    PRIORITY_DISPLAY,
    QUERY_DIMENSIONS,
    STATUS_MAPPING,
    STATUS_REVERSE_MAPPING,
    TASK_FIELDS,
)

# 后端 API 地址
BACKEND_URL = os.environ.get("BACKEND_URL", "http://localhost:3001")

# 下载目录（用于存放生成的报告文件）
DOWNLOADS_DIR = Path(__file__).parent.parent / "server" / "data" / "downloads"
DOWNLOADS_DIR.mkdir(parents=True, exist_ok=True)


# ==================== 任务查询 ====================


def call_backend_api(method: str, path: str, data: Optional[dict] = None, params: Optional[dict] = None) -> dict:
    """
    调用后端 API

    Args:
        method: HTTP 方法 (GET/POST/PUT/DELETE)
        path: API 路径 (如 /api/tasks)
        data: 请求体数据
        params: 查询参数

    Returns:
        API 响应数据
    """
    import urllib.request

    url = f"{BACKEND_URL}{path}"

    if params:
        query_string = "&".join(f"{k}={v}" for k, v in params.items() if v is not None)
        if query_string:
            url += f"?{query_string}"

    headers = {"Content-Type": "application/json"}

    if method == "GET":
        req = urllib.request.Request(url, headers=headers, method="GET")
    else:
        body = json.dumps(data).encode() if data else b""
        req = urllib.request.Request(url, data=body, headers=headers, method=method)

    try:
        with urllib.request.urlopen(req, timeout=10) as resp:
            result = json.loads(resp.read().decode())
            return {"success": True, "data": result}
    except urllib.error.HTTPError as e:
        error_body = e.read().decode()
        return {"success": False, "error": f"API 错误 {e.code}: {error_body}"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def is_overdue(due_date: Optional[str], status: str) -> bool:
    """判断任务是否逾期"""
    if not due_date:
        return False
    # 已完成的任务不算逾期
    if status == "已完成":
        return False
    try:
        # 处理不同日期格式
        if "T" in due_date:
            # ISO 格式: 2026-06-25T10:00:00
            due = datetime.fromisoformat(due_date.replace("Z", "+00:00"))
        else:
            # 简单日期格式: 2026-06-25
            due = datetime.strptime(due_date, "%Y-%m-%d")
            # 设置为当天结束时间
            due = due.replace(hour=23, minute=59, second=59)

        # 获取当前时间（使用本地时间进行比较）
        now = datetime.now()
        return due < now
    except (ValueError, TypeError) as e:
        print(f"日期解析错误: {due_date}, {e}")
        return False


# ==================== API 端点 ====================


@asynccontextmanager
async def lifespan(app: FastAPI):
    """应用生命周期管理"""
    # 启动时检查后端 API 是否可用
    try:
        result = call_backend_api("GET", "/api/columns")
        if not result["success"]:
            print(f"Warning: Backend API not available - {result['error']}")
    except Exception as e:
        print(f"Warning: Backend API check failed - {e}")
    yield


app = FastAPI(
    title="Kanban AI Service",
    description="AI 服务，提供任务数据字典和查询接口",
    version="0.1.0",
    lifespan=lifespan,
)


@app.get("/health")
async def health():
    """健康检查"""
    return {"status": "ok", "backend_url": BACKEND_URL}


@app.get("/api/ai/dictionary")
async def get_task_dictionary():
    """
    获取任务数据字典

    返回任务字段描述、可筛选维度，供 AI 理解数据结构。
    """
    return {
        "fields": TASK_FIELDS,
        "dimensions": QUERY_DIMENSIONS,
        "query_time": datetime.now(timezone.utc).isoformat(),
    }


@app.get("/api/ai/query")
async def query_tasks(
    status: Optional[str] = Query(None, description="状态筛选: 待办, 进行中, 审核, 已完成"),
    priority: Optional[str] = Query(None, description="优先级筛选: high, medium, low"),
    assignee: Optional[str] = Query(None, description="负责人筛选"),
    overdue: Optional[bool] = Query(None, description="是否逾期"),
    tags: Optional[str] = Query(None, description="标签筛选"),
    limit: int = Query(50, ge=1, le=200, description="返回数量限制"),
):
    """
    查询任务数据

    通过后端 API 查询，确保数据一致性。
    """
    # 构建查询参数
    params = {}
    if status:
        if status not in STATUS_REVERSE_MAPPING:
            raise HTTPException(400, f"不支持的状态: {status}。支持的值: {list(STATUS_REVERSE_MAPPING.keys())}")
        params["status"] = STATUS_REVERSE_MAPPING[status]
    if priority:
        if priority not in ["high", "medium", "low"]:
            raise HTTPException(400, f"不支持的优先级: {priority}。支持的值: high, medium, low")
        params["priority"] = priority
    if assignee:
        params["assignee"] = assignee

    # 调用后端 API 获取任务
    result = call_backend_api("GET", "/api/tasks", params=params if params else None)
    if not result["success"]:
        raise HTTPException(500, result["error"])

    tasks = result["data"]

    # 获取列映射
    columns_map = {}
    columns_result = call_backend_api("GET", "/api/columns")
    if columns_result["success"]:
        for col in columns_result["data"]:
            columns_map[col["id"]] = col["title"]

    # 添加状态名称和逾期标记
    for task in tasks:
        task["status"] = columns_map.get(task["columnId"], task["columnId"])
        task["overdue"] = is_overdue(task.get("dueDate"), task["status"])

    # 标签筛选（在 Python 中处理）
    if tags:
        tasks = [t for t in tasks if tags in t.get("tags", [])]

    # 逾期筛选（在 Python 中处理）
    if overdue is not None:
        tasks = [t for t in tasks if t["overdue"] == overdue]

    # 限制返回数量
    tasks = tasks[:limit]

    return {
        "total": len(tasks),
        "tasks": tasks,
        "query_time": datetime.now(timezone.utc).isoformat(),
    }


# ==================== 对话接口 ====================


class ChatRequest(BaseModel):
    """对话请求"""

    message: str
    session_id: Optional[str] = None


class ChatResponse(BaseModel):
    """对话响应"""

    content: str
    session_id: str


@app.post("/api/ai/chat", response_model=ChatResponse)
async def chat(request: ChatRequest):
    """
    AI 对话接口

    使用 Harness SDK 进行多轮对话。
    AI 会自动调用 dictionary 和 query 接口获取任务数据。
    """
    # 读取配置（支持多种环境变量格式）
    api_key = (
        os.environ.get("API_KEY")
        or os.environ.get("OPENAI_API_KEY")
        or os.environ.get("ANTHROPIC_API_KEY")
    )
    base_url = os.environ.get("API_BASE_URL") or os.environ.get("OPENAI_BASE_URL")
    model = os.environ.get("AI_MODEL", "astron-code-latest")

    if not api_key:
        return ChatResponse(
            content="AI 服务未配置 API Key。请在 .env 文件中设置 API_KEY。",
            session_id=request.session_id or "default",
        )

    # 尝试使用 Harness SDK
    try:
        from harness import AgentHarness, HarnessConfig
        from harness.memory.memory_file import MemoryScoringConfig
        from harness.tools.builtins import UpdateCoreMemoryTool

        # MEMORY.md 存放在 server/data/ 目录（与数据库同位置）
        memory_path = Path(__file__).parent.parent / "server" / "data"

        # 创建 Agent（支持第三方 OpenAI 兼容 API）
        agent = AgentHarness(
            config=HarnessConfig(
                model=model,
                provider="openai",
                api_key=api_key,
                base_url=base_url,
                memory_md_path=memory_path,
                memory_scoring=MemoryScoringConfig(
                    enable_llm_evaluation=False, # LLM 评估需要 Anthropic API，暂不启用
                    max_core_memory_tokens=3000,  # 容量限制（约 50 条记忆）
                    archive_fallback="file",      # 超限时归档到 MEMORY_ARCHIVE.md
                ),
            ),
            tools=[
                UpdateCoreMemoryTool(),
            ],
            system_prompt="""你是一个看板任务管理助手，帮助用户查询、分析和管理任务数据。

工具有：
- get_task_dictionary: 获取字段描述
- query_tasks: 查询任务（支持 status/priority/assignee/overdue 参数）
- manage_task: 管理任务（action='create'创建，action='update'更新。更新时用title匹配任务，assignee改负责人，status改状态，priority改优先级，progress改进度）
- generate_task_report: 生成 Word 报告

**操作流程**：
1. 用户要求更新/修改任务时，先调用 query_tasks 查询任务列表
2. 根据用户描述和任务列表，判断用户指的是哪个任务
3. 用任务的**实际标题**调用 manage_task 进行更新

示例：
- 用户说"任务1改负责人为王五"，先查询任务，发现标题为"1"的任务，然后用 title="1" 调用 manage_task
- 用户说"把登录功能改为进行中"，先查询任务，发现标题包含"登录"，然后用 title="登录" 调用 manage_task

回答要求：
1. 先调用工具获取数据再回答
2. 使用简洁 Markdown，不用 --- 分隔线
3. 任务列表用表格或紧凑列表
4. 重点信息用 **加粗**
5. 操作成功后，简要确认即可""",
        )

        # 注册工具
        @agent.tool(description="获取任务字段字典，了解可查询的字段和维度")
        def get_task_dictionary_tool() -> dict:
            """返回任务字段描述"""
            return {
                "fields": TASK_FIELDS,
                "dimensions": QUERY_DIMENSIONS,
            }

        @agent.tool(description="查询任务数据，支持按状态、优先级、负责人筛选。查询逾期任务时，overdue参数必须为true。")
        def query_tasks_tool(
            status: Optional[str] = None,
            priority: Optional[str] = None,
            assignee: Optional[str] = None,
            overdue: Optional[bool] = None,
        ) -> dict:
            """查询任务数据。overdue=True表示查询逾期任务，overdue=False表示查询非逾期任务。"""
            # 构建查询参数
            params = {}
            if status and status in STATUS_REVERSE_MAPPING:
                params["status"] = STATUS_REVERSE_MAPPING[status]
            if priority:
                params["priority"] = priority
            if assignee:
                params["assignee"] = assignee

            # 调用后端 API 查询任务
            result = call_backend_api("GET", "/api/tasks", params=params if params else None)

            if not result["success"]:
                return {"error": result["error"]}

            tasks = result["data"]

            # 添加状态名称映射
            columns_map = {}
            columns_result = call_backend_api("GET", "/api/columns")
            if columns_result["success"]:
                for col in columns_result["data"]:
                    columns_map[col["id"]] = col["title"]

            for task in tasks:
                task["status"] = columns_map.get(task["columnId"], task["columnId"])
                task["overdue"] = is_overdue(task.get("dueDate"), task["status"])

            # 逾期筛选（在 Python 中处理）
            if overdue is not None:
                if isinstance(overdue, str):
                    overdue = overdue.lower() == "true"
                tasks = [t for t in tasks if t["overdue"] == overdue]

            return {"total": len(tasks), "tasks": tasks}

        @agent.tool(description="管理任务。action='create'创建新任务，action='update'更新任务。title是任务标题关键词（不是ID），用于匹配任务。更新时：assignee改负责人，status改状态，priority改优先级。示例：action='update', title='登录', status='进行中'")
        def manage_task(
            action: str,
            title: Optional[str] = None,
            new_title: Optional[str] = None,
            description: Optional[str] = "",
            assignee: Optional[str] = "",
            priority: Optional[str] = "medium",
            dueDate: Optional[str] = "",
            tags: Optional[list[str]] = None,
            status: Optional[str] = "待办",
            progress: Optional[int] = None,
            progressText: Optional[str] = "",
        ) -> dict:
            """
            任务管理工具。

            Args:
                action: 操作类型 ('create' 或 'update')
                title: 任务标题关键词（不是ID）。创建时是完整标题；更新时用于模糊匹配
                new_title: 更新时的新标题（可选）
                description: 任务描述
                assignee: 负责人
                priority: 优先级 (high/medium/low)
                dueDate: 截止日期 (YYYY-MM-DD)
                tags: 标签列表
                status: 状态 (待办/进行中/审核/已完成)
                progress: 进度 (0-100)
                progressText: 进度说明

            Returns:
                操作结果
            """
            if action == "create":
                # 创建任务逻辑 - 通过后端 API
                if not title or not title.strip():
                    return {"error": "创建任务需要提供 title"}

                if priority not in ["high", "medium", "low"]:
                    return {"error": f"不支持的优先级: {priority}，必须是 high/medium/low"}

                if status not in STATUS_REVERSE_MAPPING:
                    return {"error": f"不支持的状态: {status}，必须是 待办/进行中/审核/已完成"}

                column_id = STATUS_REVERSE_MAPPING[status]

                if tags is None:
                    tags = []

                # 调用后端 API 创建任务（触发 WebSocket 广播）
                task_data = {
                    "title": title.strip(),
                    "description": description,
                    "assignee": assignee,
                    "priority": priority,
                    "dueDate": dueDate,
                    "tags": tags,
                    "columnId": column_id,
                }
                result = call_backend_api("POST", "/api/tasks", data=task_data)

                if not result["success"]:
                    return {"error": result["error"]}

                return {
                    "success": True,
                    "action": "create",
                    "task": result["data"]
                }

            elif action == "update":
                # 更新任务逻辑 - 通过标题关键词匹配
                if not title or not title.strip():
                    return {"error": "更新任务需要提供 title 关键词"}

                search_title = title.strip()

                # 通过后端 API 搜索任务
                result = call_backend_api("GET", "/api/tasks/search", params={"title": search_title})

                if not result["success"]:
                    return {"error": result["error"]}

                rows = result["data"]

                # 如果找不到，尝试去掉常见前缀再搜索
                if len(rows) == 0:
                    import re
                    cleaned_title = re.sub(r'^(任务|task|Task)\s*', '', search_title, flags=re.IGNORECASE)
                    if cleaned_title and cleaned_title != search_title:
                        result = call_backend_api("GET", "/api/tasks/search", params={"title": cleaned_title})
                        if result["success"]:
                            rows = result["data"]

                if len(rows) == 0:
                    return {"error": f"未找到标题包含 '{search_title}' 的任务"}

                # 获取列映射
                columns_result = call_backend_api("GET", "/api/columns")
                columns_map = {}
                if columns_result["success"]:
                    for col in columns_result["data"]:
                        columns_map[col["id"]] = col["title"]

                if len(rows) > 1:
                    tasks_info = []
                    for r in rows:
                        task_status = columns_map.get(r["columnId"], r["columnId"])
                        tasks_info.append(f"- {r['title']} (状态: {task_status})")
                    return {
                        "error": f"找到 {len(rows)} 个匹配的任务，请提供更精确的标题：\n" + "\n".join(tasks_info)
                    }

                # 找到唯一匹配，通过后端 API 更新
                task_id = rows[0]["id"]

                # 构建更新数据
                update_data = {}
                if new_title and new_title.strip():
                    update_data["title"] = new_title.strip()
                if description:
                    update_data["description"] = description
                if assignee:
                    update_data["assignee"] = assignee
                if priority and priority in ["high", "medium", "low"]:
                    update_data["priority"] = priority
                if dueDate:
                    update_data["dueDate"] = dueDate
                if tags is not None:
                    update_data["tags"] = tags
                if status and status in STATUS_REVERSE_MAPPING:
                    update_data["columnId"] = STATUS_REVERSE_MAPPING[status]
                if progress is not None and 0 <= progress <= 100:
                    update_data["progress"] = progress
                if progressText:
                    update_data["progressText"] = progressText

                if not update_data:
                    return {"error": "没有提供要更新的字段"}

                # 调用后端 API 更新任务
                result = call_backend_api("PUT", f"/api/tasks/{task_id}", data=update_data)

                if not result["success"]:
                    return {"error": result["error"]}

                return {
                    "success": True,
                    "action": "update",
                    "task": result["data"]
                }

            else:
                return {"error": f"不支持的操作: {action}，必须是 create/update"}

        @agent.tool(description="生成任务报告Word文档并返回下载链接。AI会根据content_hint参数生成符合需求的报告内容。返回结果包含 download_link 字段，请直接将此 Markdown 链接展示给用户。")
        def generate_task_report(
            title: str = "任务报告",
            content_hint: str = "",
            status: Optional[str] = None,
            priority: Optional[str] = None,
            assignee: Optional[str] = None,
        ) -> dict:
            """
            生成任务报告Word文档。

            Args:
                title: 报告标题
                content_hint: 内容提示，描述报告要包含什么内容（如"工作汇报，要有工作内容和进度"）
                status: 筛选状态
                priority: 筛选优先级
                assignee: 筛选负责人

            Returns:
                包含下载链接的字典
            """
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                import re

                # 1. 通过后端 API 查询任务数据
                params = {}
                if status and status in STATUS_REVERSE_MAPPING:
                    params["status"] = STATUS_REVERSE_MAPPING[status]
                if priority:
                    params["priority"] = priority
                if assignee:
                    params["assignee"] = assignee

                tasks_result = call_backend_api("GET", "/api/tasks", params=params if params else None)
                if not tasks_result["success"]:
                    return {"error": tasks_result["error"]}
                tasks = tasks_result["data"][:1000]  # 限制数量

                # 获取列映射
                columns_map = {}
                columns_result = call_backend_api("GET", "/api/columns")
                if columns_result["success"]:
                    for col in columns_result["data"]:
                        columns_map[col["id"]] = col["title"]

                # 获取评论
                for task in tasks:
                    task["status"] = columns_map.get(task["columnId"], task["columnId"])
                    comments_result = call_backend_api("GET", f"/api/tasks/{task['id']}/comments")
                    task["comments"] = comments_result["data"] if comments_result["success"] else []

                # 2. 让 AI 生成报告内容
                # 构建任务数据摘要
                tasks_summary = []
                priority_map = {"high": "高", "medium": "中", "low": "低"}
                for task in tasks:
                    task_info = {
                        "标题": task.get("title", ""),
                        "状态": task.get("status", ""),
                        "负责人": task.get("assignee", ""),
                        "优先级": priority_map.get(task.get("priority", ""), task.get("priority", "")),
                        "截止日期": task.get("dueDate", ""),
                        "进度": f"{task.get('progress', 0)}%",
                        "进度说明": task.get("progressText", ""),
                        "描述": task.get("description", ""),
                        "评论数": len(task.get("comments", [])),
                    }
                    tasks_summary.append(task_info)

                # 调用 LLM 生成报告内容
                report_prompt = f"""请根据以下任务数据生成一份报告，报告格式为 Markdown。

报告标题：{title}
用户需求：{content_hint or "标准任务报告"}

任务数据：
{json.dumps(tasks_summary, ensure_ascii=False, indent=2)}

要求：
1. 根据 "{content_hint}" 的需求组织报告内容
2. 使用 Markdown 格式
3. 标题用 ##，小标题用 ###
4. 重点内容用 **加粗**
5. 列表用 - 开头
6. 不要写"根据任务数据"之类的开场白，直接输出报告内容
7. 报告要有实质性内容，不要空洞"""

                # 使用 OpenAI API 生成报告
                import openai
                client = openai.OpenAI(
                    api_key=api_key,
                    base_url=base_url,
                )

                response = client.chat.completions.create(
                    model=model,
                    messages=[{"role": "user", "content": report_prompt}],
                    max_tokens=4000,
                )

                report_content = response.choices[0].message.content

                # 3. 创建 Word 文档
                doc = Document()

                # 标题
                heading = doc.add_heading(title, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 生成时间
                now = datetime.now()
                doc.add_paragraph(f"生成时间：{now.strftime('%Y-%m-%d %H:%M:%S')}")
                doc.add_paragraph()  # 空行

                # 4. 将 Markdown 内容写入 Word（使用 markdown + beautifulsoup4）
                import markdown
                from bs4 import BeautifulSoup, NavigableString

                # 转换 Markdown 到 HTML
                html_content = markdown.markdown(
                    report_content,
                    extensions=['fenced_code', 'tables']
                )
                soup = BeautifulSoup(html_content, 'html.parser')

                def process_element(element, doc):
                    """递归处理 HTML 元素并写入 Word"""
                    if isinstance(element, NavigableString):
                        text = str(element).strip()
                        if text:
                            return text
                        return None

                    tag = element.name

                    if tag in ['h1']:
                        text = element.get_text()
                        if text:
                            doc.add_heading(text, level=1)
                    elif tag in ['h2']:
                        text = element.get_text()
                        if text:
                            doc.add_heading(text, level=2)
                    elif tag in ['h3', 'h4']:
                        text = element.get_text()
                        if text:
                            doc.add_heading(text, level=3)
                    elif tag == 'p':
                        p = doc.add_paragraph()
                        process_inline(element, p)
                    elif tag in ['ul', 'ol']:
                        for li in element.find_all('li', recursive=False):
                            text = li.get_text()
                            doc.add_paragraph(text, style='List Bullet')
                    elif tag == 'table':
                        # 表格处理
                        rows = element.find_all('tr')
                        if rows:
                            # 获取最大列数
                            max_cols = max(len(row.find_all(['th', 'td'])) for row in rows)
                            table = doc.add_table(rows=len(rows), cols=max_cols)
                            table.style = 'Table Grid'

                            for i, row in enumerate(rows):
                                cells = row.find_all(['th', 'td'])
                                for j, cell in enumerate(cells):
                                    if j < max_cols:
                                        cell_para = table.rows[i].cells[j].paragraphs[0]
                                        # 清空默认段落
                                        cell_para.clear()
                                        # 处理单元格内的格式
                                        process_inline(cell, cell_para)
                                        # 表头加粗
                                        if cell.name == 'th':
                                            for run in cell_para.runs:
                                                run.bold = True
                    elif tag == 'blockquote':
                        text = element.get_text()
                        if text:
                            doc.add_paragraph(text, style='Quote')
                    else:
                        # 其他标签，直接处理子元素
                        for child in element.children:
                            process_element(child, doc)

                def process_inline(element, paragraph):
                    """处理行内元素（加粗、斜体等）"""
                    for child in element.children:
                        if isinstance(child, NavigableString):
                            text = str(child)
                            if text.strip():
                                paragraph.add_run(text)
                        elif child.name == 'strong' or child.name == 'b':
                            paragraph.add_run(child.get_text()).bold = True
                        elif child.name == 'em' or child.name == 'i':
                            paragraph.add_run(child.get_text()).italic = True
                        elif child.name == 'code':
                            run = paragraph.add_run(child.get_text())
                            run.font.name = 'Courier New'
                        elif child.name == 'br':
                            paragraph.add_run('\n')
                        elif child.name == 'span':
                            # 忽略 span 标签，只取文本
                            process_inline(child, paragraph)
                        else:
                            # 其他标签，递归处理
                            process_inline(child, paragraph)

                # 处理所有顶层元素
                for element in soup.children:
                    if isinstance(element, NavigableString):
                        text = str(element).strip()
                        if text:
                            doc.add_paragraph(text)
                    else:
                        process_element(element, doc)

                # 5. 保存文件
                filename = f"task_report_{now.strftime('%Y%m%d_%H%M%S')}.docx"
                filepath = DOWNLOADS_DIR / filename
                doc.save(str(filepath))

                # 6. 清理旧文件（保留最近10个）
                report_files = sorted(
                    DOWNLOADS_DIR.glob("task_report_*.docx"),
                    key=lambda f: f.stat().st_mtime,
                    reverse=True
                )
                for old_file in report_files[10:]:
                    old_file.unlink()

                # 7. 返回下载链接
                download_url = f"/downloads/{filename}"
                return {
                    "success": True,
                    "filename": filename,
                    "download_url": download_url,
                    "download_link": f"[点击下载：{filename}]({download_url})",
                    "total_tasks": len(tasks),
                }

            except ImportError as e:
                return {"error": f"依赖未安装: {str(e)}"}
            except Exception as e:
                return {"error": f"生成报告失败: {str(e)}"}

        # 运行对话
        result = await agent.run(
            request.message,
            session_id=request.session_id,
        )

        return ChatResponse(
            content=result.content,
            session_id=request.session_id or "default",
        )

    except ImportError:
        return ChatResponse(
            content="Harness SDK 未安装。请确保 SDK 路径正确。",
            session_id=request.session_id or "default",
        )
    except Exception as e:
        return ChatResponse(
            content=f"AI 服务错误: {str(e)}",
            session_id=request.session_id or "default",
        )


# ==================== 启动服务 ====================

if __name__ == "__main__":
    import uvicorn

    port = int(os.environ.get("AI_SERVICE_PORT", 3002))
    uvicorn.run(app, host="0.0.0.0", port=port)
