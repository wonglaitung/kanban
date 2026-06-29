"""
GenerateReportTool - 生成任务报告
"""

import json
from datetime import datetime
from typing import Any, Optional

from harness.tools.base import Tool, ToolContext
from harness.types import ToolResult

from .helpers import BACKEND_URL, DOWNLOADS_DIR, call_backend_api, get_columns_mapping


class GenerateReportTool(Tool):
    """生成任务报告工具"""

    # API 配置（从环境变量获取）
    _api_key = None
    _base_url = None
    _model = None

    def __init__(self, api_key: str = None, base_url: str = None, model: str = None):
        self._api_key = api_key
        self._base_url = base_url
        self._model = model

    @property
    def name(self) -> str:
        return "generate_task_report"

    @property
    def description(self) -> str:
        return (
            "生成任务报告Word文档并返回下载链接。"
            "AI会根据content_hint参数生成符合需求的报告内容。"
            "返回结果包含 download_link 字段，请直接将此 Markdown 链接展示给用户。"
        )

    @property
    def input_schema(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "title": {
                    "type": "string",
                    "description": "报告标题",
                    "default": "任务报告",
                },
                "content_hint": {
                    "type": "string",
                    "description": "内容提示，描述报告要包含什么内容（如'工作汇报，要有工作内容和进度'）",
                },
                "status": {
                    "type": "string",
                    "description": "筛选状态",
                },
                "priority": {
                    "type": "string",
                    "description": "筛选优先级",
                },
                "assignee": {
                    "type": "string",
                    "description": "筛选负责人",
                },
            },
            "required": [],
        }

    async def execute(
        self, arguments: dict[str, Any], context: ToolContext
    ) -> ToolResult:
        """生成任务报告"""
        title = arguments.get("title", "任务报告")
        content_hint = arguments.get("content_hint", "")
        status = arguments.get("status")
        priority = arguments.get("priority")
        assignee = arguments.get("assignee")

        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError as e:
            return ToolResult(
                tool_call_id="",
                success=False,
                content="",
                error=f"依赖未安装: {str(e)}",
            )

        # 1. 通过后端 API 查询任务数据
        params = {}
        title_to_id, _ = get_columns_mapping()
        if status and status in title_to_id:
            params["status"] = title_to_id[status]
        if priority:
            params["priority"] = priority
        if assignee:
            params["assignee"] = assignee

        tasks_result = call_backend_api(
            "GET", "/api/tasks", params=params if params else None
        )
        if not tasks_result["success"]:
            return ToolResult(
                tool_call_id="",
                success=False,
                content="",
                error=tasks_result["error"],
            )
        tasks = tasks_result["data"][:1000]  # 限制数量

        # 获取列映射
        columns_map = {}
        columns_result = call_backend_api("GET", "/api/columns")
        if columns_result["success"]:
            for col in columns_result["data"]:
                columns_map[col["id"]] = col["title"]

        # 获取评论
        for task in tasks:
            task["status"] = columns_map.get(task["columnId"], task["columnId"])
            comments_result = call_backend_api("GET", f"/api/tasks/{task['id']}/comments")
            task["comments"] = (
                comments_result["data"] if comments_result["success"] else []
            )

        # 2. 让 AI 生成报告内容
        tasks_summary = []
        priority_map = {"high": "高", "medium": "中", "low": "低"}
        for task in tasks:
            task_info = {
                "标题": task.get("title", ""),
                "状态": task.get("status", ""),
                "负责人": task.get("assignee", ""),
                "优先级": priority_map.get(
                    task.get("priority", ""), task.get("priority", "")
                ),
                "截止日期": task.get("dueDate", ""),
                "进度": f"{task.get('progress', 0)}%",
                "进度说明": task.get("progressText", ""),
                "描述": task.get("description", ""),
                "评论数": len(task.get("comments", [])),
            }
            tasks_summary.append(task_info)

        report_prompt = f"""请根据以下任务数据生成一份报告，报告格式为 Markdown。

报告标题：{title}
用户需求：{content_hint or "标准任务报告"}

任务数据：
{json.dumps(tasks_summary, ensure_ascii=False, indent=2)}

要求：
1. 根据 "{content_hint}" 的需求组织报告内容
2. 使用 Markdown 格式
3. 标题用 ##，小标题用 ###
4. 重点内容用 **加粗**
5. 列表用 - 开头
6. 不要写"根据任务数据"之类的开场白，直接输出报告内容
7. 报告要有实质性内容，不要空洞"""

        # 使用 OpenAI API 生成报告
        try:
            import openai

            client = openai.OpenAI(
                api_key=self._api_key,
                base_url=self._base_url,
            )

            response = client.chat.completions.create(
                model=self._model,
                messages=[{"role": "user", "content": report_prompt}],
                max_tokens=4000,
            )

            report_content = response.choices[0].message.content
        except Exception as e:
            return ToolResult(
                tool_call_id="",
                success=False,
                content="",
                error=f"生成报告内容失败: {str(e)}",
            )

        # 3. 创建 Word 文档
        doc = Document()

        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        now = datetime.now()
        doc.add_paragraph(f"生成时间：{now.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()

        # 4. 将 Markdown 内容写入 Word
        try:
            import markdown
            from bs4 import BeautifulSoup, NavigableString

            html_content = markdown.markdown(
                report_content, extensions=["fenced_code", "tables"]
            )
            soup = BeautifulSoup(html_content, "html.parser")

            self._process_html_to_doc(soup, doc)

        except ImportError:
            # 回退到简单文本
            doc.add_paragraph(report_content)

        # 5. 保存文件
        filename = f"task_report_{now.strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = DOWNLOADS_DIR / filename
        doc.save(str(filepath))

        # 设置文件权限，让 nginx 可以读取
        import os
        os.chmod(filepath, 0o644)

        # 6. 清理旧文件（保留最近10个）
        report_files = sorted(
            DOWNLOADS_DIR.glob("task_report_*.docx"),
            key=lambda f: f.stat().st_mtime,
            reverse=True,
        )
        for old_file in report_files[10:]:
            old_file.unlink()

        # 7. 返回下载链接
        download_url = f"/downloads/{filename}"
        return ToolResult(
            tool_call_id="",
            success=True,
            content=json.dumps(
                {
                    "success": True,
                    "filename": filename,
                    "download_url": download_url,
                    "download_link": f"[点击下载：{filename}]({download_url})",
                    "total_tasks": len(tasks),
                },
                ensure_ascii=False,
            ),
        )

    def _process_html_to_doc(self, soup, doc):
        """处理 HTML 并写入 Word 文档"""
        from bs4 import NavigableString

        def process_element(element, doc):
            if isinstance(element, NavigableString):
                text = str(element).strip()
                if text:
                    return text
                return None

            tag = element.name

            if tag in ["h1"]:
                text = element.get_text()
                if text:
                    doc.add_heading(text, level=1)
            elif tag in ["h2"]:
                text = element.get_text()
                if text:
                    doc.add_heading(text, level=2)
            elif tag in ["h3", "h4"]:
                text = element.get_text()
                if text:
                    doc.add_heading(text, level=3)
            elif tag == "p":
                p = doc.add_paragraph()
                self._process_inline(element, p)
            elif tag in ["ul", "ol"]:
                for li in element.find_all("li", recursive=False):
                    text = li.get_text()
                    doc.add_paragraph(text, style="List Bullet")
            elif tag == "table":
                self._process_table(element, doc)
            elif tag == "blockquote":
                text = element.get_text()
                if text:
                    doc.add_paragraph(text, style="Quote")
            else:
                for child in element.children:
                    process_element(child, doc)

        for element in soup.children:
            if isinstance(element, NavigableString):
                text = str(element).strip()
                if text:
                    doc.add_paragraph(text)
            else:
                process_element(element, doc)

    def _process_inline(self, element, paragraph):
        """处理行内元素"""
        from bs4 import NavigableString

        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text.strip():
                    paragraph.add_run(text)
            elif child.name == "strong" or child.name == "b":
                paragraph.add_run(child.get_text()).bold = True
            elif child.name == "em" or child.name == "i":
                paragraph.add_run(child.get_text()).italic = True
            elif child.name == "code":
                run = paragraph.add_run(child.get_text())
                run.font.name = "Courier New"
            elif child.name == "br":
                paragraph.add_run("\n")
            elif child.name == "span":
                self._process_inline(child, paragraph)
            else:
                self._process_inline(child, paragraph)

    def _process_table(self, element, doc):
        """处理表格"""
        rows = element.find_all("tr")
        if not rows:
            return

        max_cols = max(len(row.find_all(["th", "td"])) for row in rows)
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = "Table Grid"

        for i, row in enumerate(rows):
            cells = row.find_all(["th", "td"])
            for j, cell in enumerate(cells):
                if j < max_cols:
                    cell_para = table.rows[i].cells[j].paragraphs[0]
                    cell_para.clear()
                    self._process_inline(cell, cell_para)
                    if cell.name == "th":
                        for run in cell_para.runs:
                            run.bold = True
